#!/usr/bin/env python3
"""Generate ai-learning-roadmap-12months.docx — ML Road Map expanded to 12 months."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_rtl_paragraph(cell, text: str) -> None:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:asciiTheme"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsiTheme"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")


def add_ltr_paragraph(cell, text: str, bold: bool = False) -> None:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.bold = bold
    run.font.name = "Calibri"


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = t1.add_run("مسار تعلّم ML: من مخطط شهرين إلى ١٢ شهرًا")
    r1.bold = True
    r1.font.size = Pt(18)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = t2.add_run("ML Learning Path: 2-Month Roadmap Expanded to 12 Months")
    r2.bold = True
    r2.font.size = Pt(16)

    doc.add_paragraph(
        "Source: ML Road Map units — ~60 h/month (~720 h/year). Previous roadmaps replaced."
    )

    doc.add_heading("الوحدات والأشهر | Units → months", level=1)
    units = [
        (
            "1",
            "Python in AI",
            "مقدمة Python؛ حزم ومكتبات؛ API؛ Gradio؛ Python في علم البيانات",
            "Intro Python; packages & libraries; API; Gradio; Python in data science",
            "١–٣",
            "1–3",
        ),
        (
            "2",
            "Linear Algebra & Statistics",
            "جبر خطي؛ تحويلات و eigen؛ إحصاء وصفي؛ احتمالات وتوزيعات",
            "Linear algebra; transformations & eigen; descriptive stats; probability",
            "٤–٦",
            "4–6",
        ),
        (
            "3",
            "EDA & ML",
            "معالجة؛ انحدار؛ تصنيف؛ تجميع؛ تقليل أبعاد؛ اختيار نموذج؛ XGBoost",
            "Preprocessing; regression; classification; clustering; dim reduction; selection; XGBoost",
            "٧–٩",
            "7–9",
        ),
        (
            "4",
            "Deep Learning, CV & NLP",
            "مقدمة DL؛ CNN؛ RNN؛ رؤية حاسوب؛ مقدمة NLP",
            "Intro DL; CNN; RNN; computer vision; intro NLP",
            "١٠",
            "10",
        ),
        (
            "5",
            "GenAI & Capstone",
            "ذكاء توليدي؛ LLMs؛ وكلاء AI؛ مشروع ختامي",
            "Generative AI; LLMs; AI agents; capstone",
            "١١–١٢",
            "11–12",
        ),
    ]

    ut = doc.add_table(rows=1, cols=6)
    ut.alignment = WD_TABLE_ALIGNMENT.CENTER
    uh = ut.rows[0].cells
    hdr_u = ["Unit", "Unit name (EN)", "المحتوى (عربي)", "Topics (EN)", "أشهر", "Months"]
    for i, h in enumerate(hdr_u):
        uh[i].paragraphs[0].text = ""
        run = uh[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)
    for c in uh:
        set_cell_shading(c, "D9E2F3")

    for row in units:
        cells = ut.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = ""
            if i in (2, 3):
                add_rtl_paragraph(cells[i], text)
            elif i == 0:
                add_ltr_paragraph(cells[i], text, bold=True)
            else:
                add_ltr_paragraph(cells[i], text)

    doc.add_paragraph()

    doc.add_heading("التفصيل الشهري | Month-by-month (12)", level=1)
    months_rows = [
        ("1", "1", "مقدمة Python — صياغة، أنواع، تحكم، دوال", "Intro to Python — syntax, types, control, functions"),
        ("2", "2", "حزم ومكتبات؛ APIs — طلبات، JSON", "Packages & libraries; APIs — requests, JSON"),
        ("3", "3", "Gradio؛ Python في علم البيانات — NumPy، Pandas", "Gradio; data science — NumPy, Pandas"),
        ("4", "4", "أساسيات جبر خطي — متجهات، مصفوفات", "Linear algebra foundations — vectors, matrices"),
        ("5", "5", "تحويلات ومفاهيم eigen", "Transformations & eigen concepts"),
        ("6", "6", "إحصاء وصفي؛ احتمالات وتوزيعات", "Descriptive stats; probability & distributions"),
        ("7", "7", "معالجة بيانات؛ EDA", "Data preprocessing; EDA"),
        ("8", "8", "انحدار؛ تصنيف؛ مقاييس", "Regression; classification; metrics"),
        ("9", "9", "تجميع؛ تقليل أبعاد؛ اختيار نموذج؛ XGBoost", "Clustering; dim reduction; model selection; XGBoost"),
        ("10", "10", "مقدمة DL؛ CNN؛ RNN؛ CV؛ مقدمة NLP", "Intro DL; CNN; RNN; computer vision; intro NLP"),
        ("11", "11", "ذكاء توليدي؛ LLMs؛ مقدمة وكلاء AI", "Generative AI; LLMs; intro AI agents"),
        ("12", "12", "وكلاء AI؛ مشروع ختامي", "AI agents; capstone project"),
    ]
    mt = doc.add_table(rows=1, cols=4)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    mh = mt.rows[0].cells
    hdr_m = ["الشهر", "Month #", "محتوى (عربي)", "Content (English)"]
    for i, h in enumerate(hdr_m):
        mh[i].paragraphs[0].text = ""
        run = mh[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for c in mh:
        set_cell_shading(c, "E2EFDA")

    for ar_m, en_m, car, cen in months_rows:
        r = mt.add_row().cells
        r[0].text = ar_m
        r[1].text = en_m
        r[2].text = car
        r[3].text = cen

    doc.add_paragraph()
    doc.add_heading("ملاحظات | Notes", level=1)
    notes = doc.add_table(rows=4, cols=2)
    notes.rows[0].cells[0].text = "عربي"
    notes.rows[0].cells[1].text = "English"
    pairs = [
        ("استخدم Git ومشروعًا صغيرًا شهريًا", "Use Git + one small project per month"),
        ("الشهر ١٠ مكثف — ركّز على التطبيق العملي", "Month 10 is dense — prioritize labs"),
        ("الختام: بيانات → نموذج → عرض (مثلاً Gradio)", "Capstone: data → model → demo (e.g. Gradio)"),
    ]
    for i, (a, e) in enumerate(pairs, start=1):
        notes.rows[i].cells[0].text = a
        notes.rows[i].cells[1].text = e
    for c in notes.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True

    out = "ai-learning-roadmap-12months.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
